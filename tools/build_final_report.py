from pathlib import Path
from datetime import date
import textwrap

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from PIL import Image, ImageDraw, ImageFont

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "output" / "reports"
ASSETS = OUT / "assets"
OUT.mkdir(parents=True, exist_ok=True); ASSETS.mkdir(parents=True, exist_ok=True)
DOCX = OUT / "VentureMind_AI_Final_Project_Report.docx"

NAVY = "0B1736"; PURPLE = "5B31D6"; LAVENDER = "F1EDFF"; INK = "16213A"; MUTED = "52627B"; PALE = "F6F8FC"; GREEN = "197D53"

def diagram(path, title, columns):
    """Create a clean, self-contained raster diagram without external plotting packages."""
    width, height = 1980, 760
    image = Image.new("RGB", (width, height), "#F7F9FD")
    draw = ImageDraw.Draw(image)
    regular = ImageFont.truetype("C:/Windows/Fonts/arial.ttf", 27)
    small = ImageFont.truetype("C:/Windows/Fonts/arial.ttf", 23)
    bold = ImageFont.truetype("C:/Windows/Fonts/arialbd.ttf", 31)
    title_font = ImageFont.truetype("C:/Windows/Fonts/arialbd.ttf", 43)
    draw.text((44, 34), title, fill="#0B1736", font=title_font)
    n = len(columns); margin, gap = 42, 26
    box_w = (width - margin * 2 - gap * (n - 1)) // n
    y, box_h = 150, 490
    for i, (head, items, color) in enumerate(columns):
        x = margin + i * (box_w + gap)
        draw.rounded_rectangle((x, y, x + box_w, y + box_h), radius=18, fill="white", outline="#D6DDEC", width=3)
        draw.rounded_rectangle((x, y, x + box_w, y + 92), radius=18, fill=color)
        head_box = draw.textbbox((0, 0), head, font=bold)
        draw.text((x + (box_w - (head_box[2] - head_box[0])) / 2, y + 26), head, fill="white", font=bold)
        line_y = y + 126
        for item in items:
            wrapped = textwrap.wrap(item, width=max(13, int(box_w / 27)))
            for line in wrapped:
                draw.text((x + 25, line_y), "• " + line if line == wrapped[0] else "  " + line, fill="#25324C", font=small)
                line_y += 34
            line_y += 19
        if i < n - 1:
            cx1, cx2, cy = x + box_w + 5, x + box_w + gap - 5, y + box_h // 2
            draw.line((cx1, cy, cx2, cy), fill="#5B31D6", width=5)
            draw.polygon([(cx2, cy), (cx2 - 18, cy - 11), (cx2 - 18, cy + 11)], fill="#5B31D6")
    image.save(path)

diagram(ASSETS/"architecture.png", "VentureMind AI - Layered System Architecture", [
    ("Users", ["Founder", "Administrator", "Advisor", "Guest"], "#5B31D6"),
    ("React Frontend", ["Vite + TypeScript", "Tailwind UI", "Router / forms", "Charts"], "#315DD6"),
    ("FastAPI API", ["JWT + RBAC", "Pydantic validation", "REST endpoints", "Logging"], "#167C80"),
    ("Domain Services", ["Risk engine", "Gemini advisor", "Reports", "Advisor booking"], "#B05B26"),
    ("Persistence / Integrations", ["MySQL + Alembic", "Encrypted uploads", "Google Places optional", "Gemini optional"], "#197D53"),
])
diagram(ASSETS/"workflow.png", "Founder Journey and Decision-Support Workflow", [
    ("1. Create", ["Register / sign in", "Create startup profile", "Save project"], "#5B31D6"),
    ("2. Evaluate", ["Run deterministic risk scoring", "See explanations", "Review competitors"], "#315DD6"),
    ("3. Plan", ["Finance plan", "Requirements", "Legal checklist", "PDF report"], "#167C80"),
    ("4. Act", ["AI advisor", "Book human advisor", "Hire / operate", "Track progress"], "#197D53"),
])
diagram(ASSETS/"er.png", "Core Entity Relationship View", [
    ("Identity", ["User", "Role", "Security token", "Notification"], "#5B31D6"),
    ("Startup", ["Project", "Startup idea", "Startup profile", "Evaluation / scores"], "#315DD6"),
    ("Planning", ["Financial plan", "Milestone", "Risk assessment", "Report"], "#167C80"),
    ("People", ["Advisor profile", "Booking", "Messages", "Document requests"], "#B05B26"),
    ("Operations", ["Employee", "Task", "Attendance", "Announcement"], "#197D53"),
])

doc = Document()
sec = doc.sections[0]
sec.top_margin=Inches(0.8); sec.bottom_margin=Inches(0.75); sec.left_margin=Inches(0.85); sec.right_margin=Inches(0.85)
sec.header_distance=Inches(.35); sec.footer_distance=Inches(.35)

styles=doc.styles
normal=styles['Normal']; normal.font.name='Calibri'; normal._element.rPr.rFonts.set(qn('w:ascii'),'Calibri'); normal._element.rPr.rFonts.set(qn('w:hAnsi'),'Calibri'); normal.font.size=Pt(10.5); normal.font.color.rgb=RGBColor.from_string(INK); normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.line_spacing=1.12
for name,size,color,before,after in [('Heading 1',16,PURPLE,16,7),('Heading 2',13,'315DD6',12,5),('Heading 3',11,'0B1736',9,4)]:
    s=styles[name]; s.font.name='Calibri'; s._element.rPr.rFonts.set(qn('w:ascii'),'Calibri'); s._element.rPr.rFonts.set(qn('w:hAnsi'),'Calibri'); s.font.size=Pt(size); s.font.bold=True; s.font.color.rgb=RGBColor.from_string(color); s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after); s.paragraph_format.keep_with_next=True

def shade(cell, fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),fill); tcPr.append(shd)
def margins(cell):
    tc=cell._tc; tcPr=tc.get_or_add_tcPr(); m=tcPr.first_child_found_in('w:tcMar')
    if m is None: m=OxmlElement('w:tcMar'); tcPr.append(m)
    for side in ('top','start','bottom','end'):
        e=OxmlElement(f'w:{side}'); e.set(qn('w:w'),'110'); e.set(qn('w:type'),'dxa'); m.append(e)
def table(headers, rows, widths=None):
    t=doc.add_table(rows=1, cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.style='Table Grid'; t.autofit=False
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=h; shade(c,'EDE9FE'); margins(c); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for r in c.paragraphs[0].runs: r.font.bold=True; r.font.size=Pt(9); r.font.color.rgb=RGBColor.from_string(PURPLE)
    for row in rows:
        cells=t.add_row().cells
        for i,val in enumerate(row):
            cells[i].text=str(val); margins(cells[i]); cells[i].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after=Pt(2)
                for r in p.runs: r.font.size=Pt(8.8)
    if widths:
        for row in t.rows:
            for i,w in enumerate(widths): row.cells[i].width=Inches(w)
    doc.add_paragraph().paragraph_format.space_after=Pt(1)
    return t
def para(text='', boldlead=None):
    p=doc.add_paragraph(); p.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    if boldlead and text.startswith(boldlead):
        p.add_run(boldlead).bold=True; p.add_run(text[len(boldlead):])
    else: p.add_run(text)
    return p
def bullet(text):
    p=doc.add_paragraph(style='List Bullet'); p.paragraph_format.space_after=Pt(3); p.add_run(text); return p
def number(text):
    p=doc.add_paragraph(style='List Number'); p.paragraph_format.space_after=Pt(3); p.add_run(text); return p
def callout(label, text):
    t=doc.add_table(rows=1, cols=1); t.autofit=False; t.columns[0].width=Inches(6.5); c=t.cell(0,0); shade(c,'F4F1FF'); margins(c); p=c.paragraphs[0]; r=p.add_run(label+'  '); r.bold=True; r.font.color.rgb=RGBColor.from_string(PURPLE); p.add_run(text); doc.add_paragraph().paragraph_format.space_after=Pt(1)
def fig(path, caption):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run().add_picture(str(path),width=Inches(6.35)); c=doc.add_paragraph(caption); c.alignment=WD_ALIGN_PARAGRAPH.CENTER; c.runs[0].italic=True; c.runs[0].font.size=Pt(9); c.runs[0].font.color.rgb=RGBColor.from_string(MUTED)
def page_break(): doc.add_page_break()
def add_page_num(section):
    p=section.footer.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    r=p.add_run('VentureMind AI | '); r.font.size=Pt(8); r.font.color.rgb=RGBColor.from_string(MUTED)
    fld=OxmlElement('w:fldSimple'); fld.set(qn('w:instr'),'PAGE'); p._p.append(fld)

for s in doc.sections: add_page_num(s)

# Cover
for _ in range(5): doc.add_paragraph()
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('FINAL YEAR SOFTWARE ENGINEERING PROJECT REPORT'); r.bold=True; r.font.size=Pt(13); r.font.color.rgb=RGBColor.from_string(PURPLE)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(16); r=p.add_run('VentureMind AI'); r.bold=True; r.font.size=Pt(31); r.font.color.rgb=RGBColor.from_string(NAVY)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('AI-Powered Startup Lifecycle Management, Risk Analysis and Business Advisory Platform'); r.font.size=Pt(15); r.font.color.rgb=RGBColor.from_string(PURPLE)
doc.add_paragraph()
callout('Project purpose', 'To help entrepreneurs move from a startup idea to evidence-led planning, risk assessment, practical preparation, human guidance and early business operations through one integrated web platform.')
doc.add_paragraph()
table(['Prepared by','Programme','Module','Academic year'], [['Vaishnavadevy Vasanthakumar (st20312021)','BSc Software Engineering','CSE6035 Development Project - WRIT1','2025/2026']], [2.0,1.7,1.6,1.2])
doc.add_paragraph('Institution: International College of Business & Technology (ICBT)',style='Normal').alignment=WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('Supervisor: Mr. K. Sajuran',style='Normal').alignment=WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph(f'Report version: Current implementation review - {date.today().strftime("%d %B %Y")}',style='Normal').alignment=WD_ALIGN_PARAGRAPH.CENTER
page_break()

doc.add_heading('Executive Summary',0)
para('VentureMind AI is a full-stack, AI-assisted startup lifecycle management platform designed for aspiring entrepreneurs who need structured support before and during the early stages of launching a business. The project responds to a common practical problem: founders often make high-impact decisions about market need, investment, legal preparation and operations with incomplete evidence and disconnected tools. VentureMind AI combines a guided startup profile, deterministic explainable risk scoring, planning tools, advisory support and operational management features in one web application.')
para('The implemented system uses React with TypeScript and Tailwind CSS on the frontend, and FastAPI with SQLAlchemy, Alembic and MySQL on the backend. JWT-based authentication and role-based access control support founders, administrators, legal advisors and business mentors. The platform includes a startup workspace, risk analysis, financial planning, legal checklists, report generation, an AI business advisor, human advisor booking, hiring support, poster generation, business operations, an administrator control centre and supporting founder resources.')
para('A key design decision is that the risk engine does not produce random scores. It calculates factor scores from evidence present in the founder profile using documented weighted rules and returns positive factors, gaps and practical next actions. A compatible Gemini integration can generate contextual advisory content where an API key is available; a clearly labelled structured fallback is provided where it is not. The current implementation is designed to operate locally first and provides a scalable base for future cloud deployment, formal payment integration, verified Places search and empirical user evaluation.')
callout('Current status', 'The core frontend and FastAPI backend are implemented. The database schema is managed with Alembic migrations. Some external services such as Gemini, Google Places, payment processing and transactional email require separately configured credentials for live production use.')

doc.add_heading('Table of Contents',1)
for item in ['1. Introduction and Project Context','2. Literature and Technology Review','3. Methodology and Project Planning','4. Requirements Specification','5. System Design and Architecture','6. Implementation','7. Testing and Evaluation','8. Ethics, Security and Limitations','9. Conclusion and Future Work','References','Appendices']:
    doc.add_paragraph(item,style='List Number')

doc.add_heading('1. Introduction and Project Context',1)
doc.add_heading('1.1 Background',2)
para('Entrepreneurs frequently begin with a business idea but lack a single, structured environment in which to articulate their target customer, identify assumptions, estimate costs, consider compliance obligations and record subsequent decisions. General-purpose documents and calculators are useful but place the integration burden on the founder. VentureMind AI addresses this gap by offering a guided workflow from idea validation to operational preparation.')
doc.add_heading('1.2 Problem Statement',2)
para('Early-stage founders may overestimate demand, overlook legal or financial requirements, or delay customer validation because information is fragmented and analysis is difficult to interpret. Existing AI idea generators can produce suggestions, but they do not necessarily explain why a specific idea may carry risk or guide the user across the wider business lifecycle. The problem investigated is how a transparent software platform can provide integrated, explainable and actionable startup decision support without presenting uncertain outputs as guaranteed business outcomes.')
doc.add_heading('1.3 Aim and Objectives',2)
para('The aim is to design and develop an AI-powered startup lifecycle management system that helps founders assess, plan and manage an early business venture through explainable decision support.')
for x in ['Design a responsive role-based web application for founders, administrators and approved advisors.','Capture a structured startup profile including market, location, budget, experience and business goals.','Calculate transparent risk, confidence and success-related indicators using deterministic weighted factors.','Generate planning outputs including business requirements, financial forecasts, legal checklists and downloadable reports.','Provide contextual AI guidance and a controlled human-advisor booking workflow.','Persist key project, evaluation, booking and operations data in a normalized MySQL database.','Evaluate the system through functional testing, API validation and planned usability feedback.']: bullet(x)
doc.add_heading('1.4 Scope',2)
para('The scope covers local deployment, structured startup planning, evidence-based scoring, user and advisor workflows, reporting, administration and operational features. The platform is a decision-support tool rather than a substitute for legal, financial, investment or market-research professionals. Live map data, real payment collection and commercial email delivery are treated as integration-ready extensions rather than guaranteed local features.')

doc.add_heading('2. Literature and Technology Review',1)
doc.add_heading('2.1 Startup Validation and Lean Experimentation',2)
para('Lean startup practice emphasises testing assumptions about the problem, customer and willingness to pay before investing heavily in a solution (Ries, 2011). VentureMind operationalises this concept by turning missing evidence into clear next actions, for example conducting customer interviews, testing a price point, or identifying local competitors. The product therefore focuses on learning rather than claiming predictive certainty.')
doc.add_heading('2.2 Explainable Decision Support',2)
para('Explainable AI is important when users need to understand how a recommendation was reached. Rather than treating a numerical outcome as a black box, VentureMind associates each risk or confidence output with supporting evidence, missing information and recommended actions. This design is consistent with the objective of interpretable machine learning and accountable automated decision support (Molnar, 2022).')
doc.add_heading('2.3 Relevant Technology Choices',2)
table(['Technology','Role in VentureMind','Reason for selection'], [['React, TypeScript and Vite','Client application','Component reuse, strong typing and rapid development'],['Tailwind CSS','Design system','Consistent responsive styling with low CSS overhead'],['FastAPI and Pydantic','REST API','Typed contracts, automatic Swagger documentation and validation'],['SQLAlchemy, Alembic and MySQL','Persistence','ORM abstraction, repeatable schema migrations and relational integrity'],['Gemini-compatible LLM interface','Contextual advisory text','Optional natural-language business guidance'],['scikit-learn / SHAP-ready design','Analytics extension','Supports future trained models and explanatory features']], [1.55,2.1,2.85])
para('The platform follows clean-architecture-oriented separation: presentation components call API clients; endpoint modules validate requests; services hold domain rules; models represent persisted entities; and database sessions are injected per request. This separation supports testing and gradual replacement of optional external services.')

doc.add_heading('3. Methodology and Project Planning',1)
doc.add_heading('3.1 Development Methodology',2)
para('An iterative Agile-inspired approach was selected because the project contains multiple connected modules whose user flows require repeated refinement. Each increment produces a visible feature, followed by local testing and UI review. The approach is compatible with Design Science Research because the software artefact is constructed to address a real entrepreneurial decision-support problem, then evaluated against functional and usability criteria.')
table(['Phase','Main deliverables','Current outcome'], [['1. Foundation','Architecture, database, authentication, landing page','Implemented'],['2. Founder planning','Profile, workspace, risk analysis, finance and requirements','Implemented with deterministic fallback'],['3. Guidance','AI advisor, legal checklist, competitor support','Implemented; external data configurable'],['4. Operations','Hiring, posters, employee and task operations','Implemented'],['5. Governance','Reports, admin workspace, advisor review and bookings','Implemented; migrations must be applied'],['6. Evaluation','Functional tests, usability feedback, research analysis','Partially completed / recommended next']], [1.0,3.7,1.8])
doc.add_heading('3.2 Research and Evaluation Approach',2)
para('The final thesis should supplement technical testing with a small usability evaluation. A suggested study is to recruit 5-10 prospective founders or students, give them structured tasks such as creating a profile, interpreting a risk result and booking an advisor, and collect task completion, perceived clarity, System Usability Scale (SUS) responses and qualitative comments. No fabricated results should be reported; only data actually collected with participant consent may be included.')

doc.add_heading('4. Requirements Specification',1)
doc.add_heading('4.1 Actors',2)
table(['Actor','Primary responsibilities'], [['Guest','View landing page, rapid validation and public learning resources'],['Founder','Create profile, evaluate idea, plan finance, generate reports, request human support'],['Administrator','Manage users, content, announcements, feedback, AI monitoring and advisor verification'],['Legal Advisor / Business Mentor','Set availability, accept appointments, communicate with founders and request documents'],['Job Applicant / Investor','Reserved role support for future targeted workflows']], [1.8,4.7])
doc.add_heading('4.2 Functional Requirements',2)
for x in ['FR1: The system shall register users, authenticate with JWT tokens and restrict protected routes by role.','FR2: The system shall allow founders to create and update a structured startup profile.','FR3: The system shall compute risk and confidence-related outputs from documented inputs rather than random values.','FR4: The system shall present explanations, improvement suggestions and planning outputs.','FR5: The system shall generate financial plans, requirements checklists and Sri Lanka-oriented legal guidance.','FR6: The system shall create downloadable business reports from stored project/evaluation data.','FR7: The system shall support AI advisory chat with saved startup context and a transparent fallback.','FR8: The system shall enable founders to book approved advisors and view status updates.','FR9: The system shall enable advisors to manage availability, messages, meeting details and requested documents.','FR10: The system shall provide administrators with content, account, feedback, analytics and audit capabilities.']: bullet(x)
doc.add_heading('4.3 Non-functional Requirements',2)
table(['Category','Requirement'], [['Usability','Responsive layout, consistent navigation, loading, empty and error states'],['Security','JWT authentication, RBAC, validation, private encrypted advisor uploads and no public document URLs'],['Reliability','Database migrations, request-scoped sessions, logging and consistent API response envelopes'],['Maintainability','Type hints, modular services, reusable React components and documented environment settings'],['Performance','Client-side route rendering, bounded API lists and database indexes for frequently filtered entities'],['Ethics','Decision-support disclaimers, consent before identity-document collection and no fabricated evidence']], [1.7,4.8])

doc.add_heading('5. System Design and Architecture',1)
doc.add_heading('5.1 High-Level Architecture',2)
fig(ASSETS/'architecture.png','Figure 1. VentureMind AI layered architecture.')
para('The browser client communicates with FastAPI through REST endpoints under /api/v1. FastAPI validates payloads using Pydantic schemas and delegates business rules to service modules. SQLAlchemy maps domain models to MySQL tables, while Alembic records ordered database changes. Optional integrations are isolated so the core planning experience can still run locally when a third-party key is unavailable.')
doc.add_heading('5.2 Founder Workflow',2)
fig(ASSETS/'workflow.png','Figure 2. Founder workflow from startup profile to action.')
doc.add_heading('5.3 Database Design',2)
fig(ASSETS/'er.png','Figure 3. Core entity relationship view.')
para('The database is normalized around the User and Project entities. A user can own multiple projects; a project has startup ideas, profiles, evaluations, scores, reports and lifecycle planning data. Advisor workflows use separate profiles, booking requests, availability slots, messages, document requests and payment records. This avoids storing operational, identity and business-analysis data in a single oversized table.')
table(['Area','Important tables'], [['Identity and access','users, security_tokens, notifications'],['Startup lifecycle','projects, startup_ideas, startup_profiles, evaluations, evaluation_scores, reports'],['Planning','lifecycle_risk_assessments, lifecycle_financial_plans, lifecycle_milestones'],['Advisor services','advisor_profiles, advisor_booking_requests, advisor_availability_slots, advisor_booking_messages, advisor_document_requests'],['Operations','employees, attendance_records, leave_requests, operation_tasks, announcements'],['Administration','feedback, content_items, audit_logs, platform_announcements']], [1.8,4.7])
doc.add_heading('5.4 API Design',2)
table(['API group','Examples of capabilities'], [['Authentication','/auth/register, /auth/login, password reset and verification workflows'],['Projects and evaluation','/projects, /evaluations, dashboard snapshot and report generation'],['Lifecycle planning','profile, risk analysis, finance and requirements endpoints'],['Human advisors','directory, booking, slots, messages, document request/upload endpoints'],['Administration','analytics, user management, feedback, content, audit and advisor-review endpoints']], [1.8,4.7])
doc.add_heading('5.5 Explainable Evaluation Design',2)
para('The evaluation engine begins with a baseline per dimension and applies bounded evidence adjustments. Evidence that directly supports a factor improves confidence or reduces the associated risk; missing evidence increases uncertainty. The score is clamped to 0-100. Each dimension returns a numeric value together with positive factors, negative factors, rationale and next-step recommendations. The overall score is a documented weighted aggregation of dimensions, not an opaque LLM-generated number.')
table(['Dimension','Illustrative evidence considered','Output behaviour'], [['Customer / market risk','Target customer, problem clarity, demand validation, location','Higher risk when customer and demand evidence are missing'],['Financial feasibility','Budget, revenue assumptions, monthly costs, break-even inputs','Highlights insufficient runway or untested pricing'],['Legal / operational risk','Country, sector requirements, registrations and operating plan','Prompts local registration and compliance checks'],['Scalability / technical risk','Team capability, technology needs, repeatable process','Highlights delivery constraints and process gaps']], [1.45,3.25,1.8])

doc.add_heading('6. Implementation',1)
doc.add_heading('6.1 Founder Experience',2)
para('The landing page introduces explainable AI and rapid validation. The founder dashboard uses a numbered journey so the user can identify the next relevant action. The startup workspace captures business identity, category, market, target customers, location, budget and goals in progressive steps. It then provides cards for competitor analysis, risk analysis, finance, business requirements, legal guidance and advisory support.')
doc.add_heading('6.2 Administration and Governance',2)
para('The administrator control centre is intentionally separate from the founder dashboard. It provides platform metrics, announcements, feedback review, user management, published-content management, audit logs and advisor application review. The advisor workflow is designed so a professional first applies through the common frontend; an administrator verifies the submission and approves or rejects it. Approved professionals are then surfaced in the founder directory.')
doc.add_heading('6.3 Advisor and Appointment Management',2)
para('Founders can browse approved Legal Advisors and Business Mentors, select an appointment type and optionally reserve an available time slot. Advisors can manage availability, accept or decline requests, provide a meeting link, send a response, request specific documents and review files shared by the founder. Document metadata is stored in MySQL while file bytes are encrypted in private storage. Demonstration payment records are clearly labelled and do not charge money until a real provider is integrated.')
doc.add_heading('6.4 External Integration Boundaries',2)
for x in ['Gemini: used only where GEMINI_API_KEY is configured; otherwise the system returns structured planning guidance.','Google Places: intended for live nearby competitor searches when GOOGLE_PLACES_API_KEY and billing are configured; a Maps link remains a fallback.','Email: password reset and verification can later be connected to Brevo, Resend or SendGrid.','Payments: advisor payment records are currently demonstration-only; a compliant gateway must replace this before accepting real payments.']: bullet(x)

doc.add_heading('7. Testing and Evaluation',1)
doc.add_heading('7.1 Current Verification',2)
table(['Test area','Method','Status / evidence'], [['Frontend static quality','ESLint with zero warnings','Passed in the current workspace'],['Frontend build','TypeScript build and Vite production bundle','Passed in the current workspace'],['API contracts','FastAPI request validation and Swagger documentation','Implemented; run locally at /docs'],['Database migration','Alembic ordered revisions','Required before use; use alembic upgrade head'],['Role workflow','Protected routes and server-side role checks','Implemented; requires role-specific test accounts'],['External services','Configured-key testing','Environment dependent; not considered live without credentials']], [1.5,2.4,2.6])
doc.add_heading('7.2 Recommended Acceptance Test Cases',2)
for x in ['Register a founder; log in; create a startup profile; refresh the browser and confirm the profile persists.','Run risk analysis; verify each dimension has a score, explanation, evidence and next action.','Generate a financial plan and confirm values are retained after page refresh.','Create an advisor application, upload a permitted file and confirm that an administrator can review it.','Approve the advisor; sign out/in as that account; verify advisor dashboard access and founder directory visibility.','Create a booking, accept it as the advisor, add a meeting link and verify founder notification/status.','Generate a PDF report and confirm it contains project-specific identifiers and results.']: number(x)
doc.add_heading('7.3 Evaluation Criteria',2)
para('The system should be judged by task completion, clarity of explanation, perceived usefulness, consistency of stored data, performance of key screens and the ability to trace sensitive actions. A future empirical evaluation should compare user understanding before and after seeing explainable outputs, rather than claiming that the platform predicts commercial success.')

doc.add_heading('8. Ethics, Security and Limitations',1)
doc.add_heading('8.1 Ethical Considerations',2)
para('VentureMind AI must not misrepresent estimates as professional advice or a prediction of business success. The interface therefore uses decision-support language and recommends customer validation. Users must consent before identity-related advisor-verification data is collected. The final dissertation should document participant consent and anonymisation procedures if usability feedback is collected.')
doc.add_heading('8.2 Security Controls',2)
for x in ['Passwords are stored as hashes rather than plaintext.','JWT tokens protect authenticated API requests, while role checks restrict administrator and advisor functions.','Pydantic validates request structures and type/range constraints.','Advisor verification and founder-shared documents are intended for encrypted private storage; no public URL is persisted.','Audit logs record sensitive administrator actions, while account archival is safer than destructive deletion for important records.','Secrets such as Gemini and Places keys must remain only in backend environment files and never be committed to source control.']: bullet(x)
doc.add_heading('8.3 Current Limitations',2)
for x in ['Risk scoring is rules-based and evidence-dependent; it is not trained on longitudinal startup success data.','Competitor data requires a live provider configuration to return real listings in application screens.','Production-grade payment, email, calendar and meeting integrations are not yet enabled.','Usability findings must be collected from real participants before reporting numerical claims.','Local WAMP/MySQL environments require all Alembic migrations to be applied consistently; partial manual schema changes can cause version mismatch errors.']: bullet(x)

doc.add_heading('9. Conclusion and Future Work',1)
para('VentureMind AI demonstrates a coherent approach to building a startup decision-support platform rather than a standalone idea generator. The project integrates structured data collection, deterministic explainable scoring, planning tools, role-based advisory workflows and operational support into one product. Its principal contribution is the connection between a founder’s saved startup context and a sequence of practical next steps.')
para('The most valuable next step is empirical evaluation with real target users, followed by controlled integration of live competitor search, email, calendar and payment services. Future research may compare the rule-based explainable approach with an appropriately trained predictive model, while retaining transparency, fairness checks and user control.')

doc.add_heading('References',1)
for ref in ['FastAPI (2026) FastAPI Documentation. Available at: https://fastapi.tiangolo.com/ (Accessed: 18 August 2026).','Google (2026) Google Maps Platform: Places API Documentation. Available at: https://developers.google.com/maps/documentation/places (Accessed: 18 August 2026).','Molnar, C. (2022) Interpretable Machine Learning, 2nd edn. Available at: https://christophm.github.io/interpretable-ml-book/ (Accessed: 18 August 2026).','OWASP Foundation (2021) OWASP Application Security Verification Standard. Available at: https://owasp.org/www-project-application-security-verification-standard/ (Accessed: 18 August 2026).','Pydantic (2026) Pydantic Documentation. Available at: https://docs.pydantic.dev/ (Accessed: 18 August 2026).','Ries, E. (2011) The Lean Startup. New York: Crown Business.','SQLAlchemy (2026) SQLAlchemy Documentation. Available at: https://docs.sqlalchemy.org/ (Accessed: 18 August 2026).','The React Team (2026) React Documentation. Available at: https://react.dev/ (Accessed: 18 August 2026).']: doc.add_paragraph(ref, style='List Bullet')

doc.add_heading('Appendix A: Local Installation and Execution',1)
for command in ['cd "C:\\Users\\USER\\OneDrive\\Documents\\Venture\\backend"',' .\\.venv\\Scripts\\Activate.ps1','python -m alembic upgrade head','python -m uvicorn app.main:app --host 127.0.0.1 --port 8000 --reload','','cd "C:\\Users\\USER\\OneDrive\\Documents\\Venture\\frontend"','npm run dev']:
    p=doc.add_paragraph(); r=p.add_run(command or ''); r.font.name='Consolas'; r.font.size=Pt(9); r.font.color.rgb=RGBColor.from_string('1F3A5F')
doc.add_heading('Appendix B: Viva Demonstration Route',1)
for x in ['Open the landing page and explain the project objective and Explainable AI principle.','Use Rapid Validation as a guest or register a founder account.','Create a startup profile and show how profile evidence feeds the risk analysis.','Show risk explanations, financial planning, legal checklist and report generation.','Show the human advisor directory, application/approval process and advisor dashboard.','Show the admin dashboard for user, feedback, content, announcement and audit management.','Explain external services as configurable integrations, not fabricated live results.']: number(x)
doc.add_heading('Appendix C: Submission Checklist',1)
for x in ['Update title page details, supervisor name and final submission date.','Replace or extend the literature review with sources personally read and cited using Harvard style.','Add real usability-study method, consent material, participant data and results only after conducting the study.','Include signed project log sheets and supervisor approvals where required.','Run all migrations, test the local system and capture your own final screenshots.','Export the final reviewed DOCX to PDF before Turnitin submission.']: bullet(x)

doc.core_properties.title='VentureMind AI Final Project Report'; doc.core_properties.author='Vaishnavadevy Vasanthakumar'; doc.core_properties.subject='BSc Software Engineering Development Project'
doc.save(DOCX)
print(DOCX)
