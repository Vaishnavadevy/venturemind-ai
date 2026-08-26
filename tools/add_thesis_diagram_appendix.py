from pathlib import Path
from docx import Document
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[1]
PATH = ROOT / "output" / "reports" / "VentureMind_AI_Final_Project_Report.docx"
doc = Document(PATH)

def add_after(anchor, text, style=None):
    p = doc.add_paragraph(text, style=style)
    p.paragraph_format.space_after = Pt(7)
    anchor._p.addnext(p._p)
    return p

# Add literature entries directly before Appendix A so references remain together.
appendix_a = next(p for p in doc.paragraphs if p.text.strip() == "Appendix A: Local Installation and Execution")
anchor = appendix_a
for reference in reversed([
    "Creswell, J.W. and Creswell, J.D. (2018) Research Design: Qualitative, Quantitative, and Mixed Methods Approaches. 5th edn. Thousand Oaks, CA: Sage.",
    "Lundberg, S.M. and Lee, S.-I. (2017) ‘A unified approach to interpreting model predictions’, Advances in Neural Information Processing Systems, 30, pp. 4765–4774.",
    "Osterwalder, A. and Pigneur, Y. (2010) Business Model Generation. Hoboken, NJ: Wiley.",
    "Schwaber, K. and Sutherland, J. (2020) The Scrum Guide. Available at: https://scrumguides.org/ (Accessed: 18 August 2026).",
]):
    anchor = add_after(anchor, reference, "List Bullet")

doc.add_page_break()
doc.add_heading("Appendix D: Diagram Generator Prompts", 1)
doc.add_paragraph("Use the following prompts in a diagram generator, then review every label for consistency with the implemented system. Use a white background, navy text, purple accents, flat academic vector style, 16:9 landscape, high-resolution SVG or PNG, and do not include a vendor logo.")

prompts = [
    ("Figure D1 — System architecture", "Create a clean layered software architecture diagram for 'VentureMind AI', an AI-powered startup lifecycle management platform. Show four user roles on the left: Guest, Startup Founder, Administrator, Human Advisor. In the centre show React + Vite + TypeScript frontend with Tailwind CSS, React Router, React Hook Form, Axios and Recharts. Arrow to FastAPI backend with JWT authentication, RBAC, Pydantic validation, REST API, logging and Swagger. Arrow to domain services: Explainable Risk Engine, Financial Planner, PDF Report Generator, AI Business Advisor, Advisor Booking Service, Notification Service. On the right show MySQL + SQLAlchemy + Alembic and optional integrations Gemini API, Google Places API, Email/Calendar/Payment providers. Label optional integrations clearly. Use bidirectional data-flow arrows where appropriate."),
    ("Figure D2 — Founder activity workflow", "Create a UML activity diagram for a VentureMind AI startup founder. Start with Guest visits landing page, then decision Register or Rapid Validation. After login: Create project, complete Startup Profile step 1 business identity, step 2 customer and market, step 3 finance and operations, save draft. Decision profile complete? If no return to missing required step. If yes run explainable risk analysis, view score explanations and evidence gaps, choose next action: financial plan, business requirements, Sri Lanka legal checklist, generate PDF report, request human advisor. End with dashboard showing next action and notification status. Use swim lanes for Founder, Frontend, Backend/AI Service and Human Advisor."),
    ("Figure D3 — Use case diagram", "Create a UML use case diagram for VentureMind AI. Actors: Guest, Founder, Administrator, Legal Advisor, Business Mentor. Guest uses Browse landing content, Rapid validation, Register and Login. Founder uses Manage project/profile, Run risk analysis, View explanations, Generate financial plan, Download report, Browse advisors, Book appointment, Share requested documents, View notifications. Administrator uses Manage users/roles, Review advisor applications, Manage content/announcements, View analytics, Review feedback, Audit actions. Advisor uses Manage profile/availability, Accept/decline/reschedule booking, Message founder, Request documents, Complete consultation. Show include relationships for authenticated access and notifications."),
    ("Figure D4 — ER diagram", "Create a crow's-foot ER diagram for VentureMind AI. Core entities: Users, Projects, StartupProfiles, Evaluations, EvaluationScores, LifecycleRiskAssessments, LifecycleFinancialPlans, Reports, Notifications. Organisation entities: Organisations, OrganisationMembers, Employees, OperationTasks, AttendanceRecords, LeaveRequests, Announcements. Advisor entities: AdvisorProfiles, AdvisorAvailabilitySlots, AdvisorBookingRequests, AdvisorBookingMessages, AdvisorDocumentRequests, AdvisorSharedDocuments, AdvisorBookingReminders, AdvisorPayments. Show primary keys as UUID id, foreign keys, and cardinalities: one User owns many Projects; one Project has one StartupProfile and many Evaluations; one Evaluation has many EvaluationScores; one AdvisorProfile belongs to one User and has many slots/bookings; one booking has many messages, document requests, shared documents and reminders. Use InnoDB/foreign-key notation."),
    ("Figure D5 — Explainable risk scoring flow", "Create a transparent decision-flow diagram titled Explainable Startup Risk Evaluation. Input block: founder profile fields including business category, target customers, location, problem, solution, budget, experience, business model, revenue model and timeline. Processing block: validation and feature extraction. Then five rule cards: Customer Risk, Market Risk, Financial Risk, Legal Risk, Scalability Risk. Each rule card uses positive evidence, missing evidence and weighted adjustments, bounded 0 to 100. Combine into Overall Risk, Success Confidence and Priority level. Output panel shows score, plain-language reasoning, positive factors, negative factors and next action. Add a note: deterministic decision support, not a guarantee of business success."),
    ("Figure D6 — Advisor booking sequence", "Create a UML sequence diagram for VentureMind AI advisor booking. Lifelines: Founder, React Frontend, FastAPI API, MySQL Database, Advisor, Notification Service. Founder selects approved advisor and availability slot; frontend submits booking; API validates JWT and slot; database creates pending booking; notification service notifies advisor; advisor accepts or reschedules; API updates booking and meeting link; notification service notifies founder; founder uploads requested protected document; API stores encrypted metadata; advisor reviews and sends message; booking completes. Include alternative branches for decline and reschedule."),
    ("Figure D7 — Data flow diagram", "Create a Level 0 data flow diagram for VentureMind AI. External entities: Guest/Founder, Administrator, Human Advisor, Optional AI/Maps/Email providers. Main processes: 1 Authentication and Access Control, 2 Startup Planning and Evaluation, 3 Reports and Recommendations, 4 Advisor Consultation Management, 5 Administration and Monitoring. Data stores: User Database, Startup Project Database, Evaluations and Reports, Advisor and Booking Data, Audit and Notification Data. Show only valid data flows, including user inputs, result views, role approvals, booking updates and optional provider requests."),
    ("Figure D8 — Deployment diagram", "Create a UML deployment diagram for local-first VentureMind AI. Developer laptop hosts a browser, Vite React frontend on localhost port 5173, FastAPI backend on localhost port 8000, and WAMP MySQL/phpMyAdmin. Show backend environment variables for JWT secret, database URL, Gemini key and Google Places key. Show optional cloud deployment future: HTTPS reverse proxy, Docker containers, managed MySQL, private object storage, email/calendar/payment providers. Clearly separate current local prototype from future production deployment with dashed boundaries."),
    ("Figure D9 — Administration workflow", "Create a professional workflow diagram for advisor verification governance. Start: applicant registers normal user account. Applicant submits advisor application and protected evidence. Administrator receives review notification, checks qualification/registration reference and expiry. Decision: approved? If no, reject or request clarification and notify applicant. If yes, assign Legal Advisor or Business Mentor role, create visible advisor profile, record audit event. Founder can then browse only approved advisors and submit booking. Admin can later edit, suspend, reactivate or archive profile. Include privacy, encryption and retention-control annotations."),
    ("Figure D10 — Project roadmap", "Create a Gantt-style roadmap for VentureMind AI final-year project: Phase 1 architecture/authentication/database; Phase 2 founder profile and explainable risk analysis; Phase 3 financial planning, requirements and legal guidance; Phase 4 AI advisor, competitor integration and reports; Phase 5 hiring, posters and business operations; Phase 6 admin dashboard and governance; Phase 7 human advisor booking and notifications; Phase 8 testing, evaluation, documentation and deployment preparation. Use academic project timeline style, purple milestone markers and a clear 'future work' lane for payments, email, calendar, live Maps and production hosting.")
]
for title, prompt in prompts:
    doc.add_heading(title, 2)
    p = doc.add_paragraph(prompt)
    p.style = doc.styles["Normal"]

doc.save(PATH)
print(PATH)
