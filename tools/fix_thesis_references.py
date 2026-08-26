from pathlib import Path
from docx import Document

ROOT = Path(__file__).resolve().parents[1]
PATH = ROOT / "output" / "reports" / "VentureMind_AI_Final_Project_Report.docx"
refs = [
    "Creswell, J.W. and Creswell, J.D. (2018) Research Design: Qualitative, Quantitative, and Mixed Methods Approaches. 5th edn. Thousand Oaks, CA: Sage.",
    "Lundberg, S.M. and Lee, S.-I. (2017) ‘A unified approach to interpreting model predictions’, Advances in Neural Information Processing Systems, 30, pp. 4765–4774.",
    "Osterwalder, A. and Pigneur, Y. (2010) Business Model Generation. Hoboken, NJ: Wiley.",
    "Schwaber, K. and Sutherland, J. (2020) The Scrum Guide. Available at: https://scrumguides.org/ (Accessed: 18 August 2026).",
]
doc = Document(PATH)
appendix_a = next(p for p in doc.paragraphs if p.text.strip() == "Appendix A: Local Installation and Execution")
for p in list(doc.paragraphs):
    if p.text in refs:
        p._element.getparent().remove(p._element)
for text in refs:
    p = doc.add_paragraph(text, style="List Bullet")
    appendix_a._p.addprevious(p._p)
doc.save(PATH)
print(PATH)
