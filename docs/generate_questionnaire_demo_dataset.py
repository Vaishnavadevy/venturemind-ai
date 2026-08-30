from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = "docs/VentureMind_Questionnaire_and_Synthetic_Demo_Data.docx"

PURPLE = "5B2BD8"
NAVY = "11203A"
PALE = "F4F0FF"
CAUTION = "FFF4D6"
GREY = "F3F5F7"

rows = [
    (1,"Student","Agree","Agree","Agree","Strongly agree","Agree","Agree","Risk analysis","Yes"),
    (2,"Startup founder","Strongly agree","Agree","Strongly agree","Agree","Strongly agree","Agree","Financial planner","Yes"),
    (3,"Student","Agree","Neutral","Agree","Agree","Agree","Neutral","Registration guide","Maybe"),
    (4,"Small-business owner","Strongly agree","Agree","Agree","Strongly agree","Agree","Agree","Risk analysis","Yes"),
    (5,"Student","Agree","Agree","Neutral","Agree","Agree","Agree","AI guidance","Maybe"),
    (6,"Startup founder","Strongly agree","Strongly agree","Agree","Strongly agree","Strongly agree","Agree","Financial planner","Yes"),
    (7,"Student","Agree","Agree","Agree","Agree","Agree","Agree","Startup profile","Yes"),
    (8,"Small-business owner","Agree","Agree","Strongly agree","Strongly agree","Agree","Agree","Risk analysis","Yes"),
    (9,"Student","Neutral","Agree","Agree","Agree","Agree","Neutral","Registration guide","Maybe"),
    (10,"Startup founder","Strongly agree","Agree","Strongly agree","Agree","Strongly agree","Strongly agree","AI guidance","Yes"),
    (11,"Student","Agree","Agree","Agree","Agree","Agree","Agree","Risk analysis","Yes"),
    (12,"Other","Agree","Neutral","Neutral","Agree","Neutral","Neutral","Launch tools","Maybe"),
    (13,"Student","Strongly agree","Agree","Agree","Strongly agree","Agree","Agree","Financial planner","Yes"),
    (14,"Small-business owner","Agree","Agree","Agree","Agree","Strongly agree","Agree","Registration guide","Yes"),
    (15,"Student","Agree","Agree","Agree","Agree","Agree","Agree","Startup profile","Yes"),
    (16,"Startup founder","Strongly agree","Strongly agree","Strongly agree","Strongly agree","Agree","Agree","Risk analysis","Yes"),
    (17,"Student","Agree","Agree","Agree","Agree","Agree","Neutral","AI guidance","Maybe"),
    (18,"Small-business owner","Agree","Neutral","Agree","Agree","Agree","Agree","Financial planner","Yes"),
    (19,"Student","Agree","Agree","Neutral","Agree","Agree","Agree","Launch tools","Maybe"),
    (20,"Startup founder","Strongly agree","Agree","Strongly agree","Strongly agree","Strongly agree","Agree","Risk analysis","Yes"),
    (21,"Student","Agree","Agree","Agree","Agree","Agree","Agree","Registration guide","Yes"),
    (22,"Small-business owner","Agree","Agree","Agree","Strongly agree","Agree","Neutral","Financial planner","Yes"),
    (23,"Student","Neutral","Neutral","Agree","Agree","Agree","Neutral","Startup profile","Maybe"),
    (24,"Startup founder","Strongly agree","Agree","Agree","Strongly agree","Strongly agree","Agree","Registration guide","Yes"),
    (25,"Student","Agree","Agree","Agree","Agree","Agree","Agree","Risk analysis","Yes"),
    (26,"Small-business owner","Agree","Agree","Strongly agree","Agree","Agree","Agree","AI guidance","Yes"),
    (27,"Student","Agree","Agree","Agree","Agree","Agree","Neutral","Launch tools","Maybe"),
    (28,"Startup founder","Strongly agree","Agree","Strongly agree","Strongly agree","Agree","Strongly agree","Financial planner","Yes"),
    (29,"Student","Agree","Agree","Agree","Agree","Agree","Agree","Registration guide","Yes"),
    (30,"Other","Neutral","Disagree","Neutral","Agree","Neutral","Neutral","Startup profile","No"),
]

def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)

def cell_text(cell, text, bold=False, color=None, size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(str(text))
    r.bold = bold
    r.font.name = "Calibri"
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def set_cell_margins(cell, top=80, start=100, bottom=80, end=100):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement('w:tcMar')
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v)); node.set(qn("w:type"), "dxa")

def format_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for row in table.rows:
        for cell in row.cells:
            set_cell_margins(cell)

def heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 9)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.bold = True; r.font.name = "Calibri"; r.font.size = Pt(16 if level == 1 else 12)
    r.font.color.rgb = RGBColor.from_string(PURPLE if level == 1 else NAVY)
    return p

def note(doc, title, text, color=CAUTION):
    table = doc.add_table(rows=1, cols=1)
    format_table(table)
    cell = table.cell(0,0); shade(cell, color)
    cell.text = ""
    p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title); r.bold = True; r.font.color.rgb = RGBColor.from_string(NAVY)
    p2 = cell.add_paragraph(text); p2.paragraph_format.space_after = Pt(0)
    return table

def question(doc, n, text, options=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"{n}. {text}"); r.bold = True; r.font.size = Pt(10.5)
    if options:
        op = doc.add_paragraph()
        op.paragraph_format.left_indent = Inches(.22); op.paragraph_format.space_after = Pt(4)
        op.add_run("☐ " + "     ☐ ".join(options)).font.size = Pt(10)

doc = Document()
section = doc.sections[0]
section.top_margin = Inches(.7); section.bottom_margin = Inches(.7)
section.left_margin = Inches(.75); section.right_margin = Inches(.75)

styles = doc.styles
styles['Normal'].font.name = 'Calibri'; styles['Normal'].font.size = Pt(10.5)
styles['Normal'].paragraph_format.space_after = Pt(6)

header = section.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
rh = header.add_run("VENTUREMIND AI | QUESTIONNAIRE TEMPLATE")
rh.font.size = Pt(8); rh.font.color.rgb = RGBColor.from_string(PURPLE); rh.bold = True
footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
rf = footer.add_run("Synthetic data appendix: demonstration only; not real research responses.")
rf.font.size = Pt(8); rf.font.color.rgb = RGBColor.from_string("666666")

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("VentureMind AI\nPrototype Usability Questionnaire")
r.bold = True; r.font.name = "Calibri"; r.font.size = Pt(22); r.font.color.rgb = RGBColor.from_string(NAVY)
p.paragraph_format.space_after = Pt(3)
p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("Ready-to-share participant form and chart-testing dataset")
r2.italic = True; r2.font.size = Pt(11); r2.font.color.rgb = RGBColor.from_string("556070")

note(doc, "Important academic use note", "Part A is a real questionnaire you may share with participants. Part B contains synthetic demonstration records only for testing Excel/Word charts and must never be described as genuine participant data, survey findings, or research validation.")

heading(doc, "Part A: Participant Information and Consent")
doc.add_paragraph("Purpose: This short questionnaire evaluates the usability and perceived usefulness of the VentureMind AI prototype. Participation is voluntary. Do not provide sensitive personal information. Responses should be used only with the participant's permission.")
question(doc, "A1", "I confirm that I understand the purpose of this questionnaire and agree to provide feedback.", ["Yes", "No"])
question(doc, "A2", "Participant role", ["Student", "Startup founder", "Small-business owner", "Other"])

heading(doc, "Part A: Questionnaire Questions")
doc.add_paragraph("For Questions 1-6, please tick one response: Strongly disagree, Disagree, Neutral, Agree, or Strongly agree.")
scale = ["Strongly disagree", "Disagree", "Neutral", "Agree", "Strongly agree"]
question(doc, 1, "The VentureMind AI dashboard is easy to understand.", scale)
question(doc, 2, "The startup-profile form is easy to complete.", scale)
question(doc, 3, "The risk-analysis explanation is understandable.", scale)
question(doc, 4, "The financial planner is useful for early business planning.", scale)
question(doc, 5, "The Sri Lanka registration guide is useful.", scale)
question(doc, 6, "The AI guidance is relevant to startup-planning questions.", scale)
question(doc, 7, "Which VentureMind feature is most useful to you?", ["Startup profile", "Risk analysis", "Financial planner", "Registration guide", "AI guidance", "Launch tools"])
question(doc, 8, "Would you consider using VentureMind AI for startup planning?", ["Yes", "Maybe", "No"])
question(doc, 9, "What is one improvement you would recommend?", None)
doc.add_paragraph("Response: _________________________________________________________________\n___________________________________________________________________________")

heading(doc, "Part B: Synthetic Demonstration Dataset (30 records)")
note(doc, "Do not report as real data", "These 30 rows are artificial examples for practicing chart creation, application demonstrations, and testing the report layout. If you use them in the dissertation, label every related table and chart as 'illustrative' or 'synthetic demonstration'.")

doc.add_paragraph("Column guide: Q1 Dashboard; Q2 Profile; Q3 Risk; Q4 Finance; Q5 Registration; Q6 AI Guidance; Q7 Most Useful Feature; Q8 Intended Use.")
headers = ["ID","Role","Q1","Q2","Q3","Q4","Q5","Q6","Q7 Feature","Q8 Use"]
table = doc.add_table(rows=1, cols=len(headers))
format_table(table)
for c, h in zip(table.rows[0].cells, headers):
    shade(c, PURPLE); cell_text(c, h, bold=True, color="FFFFFF", size=8)
for row in rows:
    cells = table.add_row().cells
    for c, value in zip(cells, row):
        cell_text(c, value, size=7.3)

heading(doc, "Chart-ready summary of the synthetic data")
summary = [
    ("Participant role", "Student 15; Startup founder 7; Small-business owner 6; Other 2", "Pie chart"),
    ("Most useful feature", "Risk analysis 7; Financial planner 6; Registration guide 6; AI guidance 4; Startup profile 4; Launch tools 3", "Pie chart"),
    ("Intended use", "Yes 20; Maybe 9; No 1", "Pie chart"),
    ("Usability ratings", "Use Q1-Q6 to calculate average Likert ratings (1-5)", "Clustered bar chart"),
]
st = doc.add_table(rows=1, cols=3); format_table(st)
for c, h in zip(st.rows[0].cells, ["Chart topic","Synthetic values","Recommended chart"]):
    shade(c, GREY); cell_text(c,h,bold=True,color=NAVY)
for a,b,cval in summary:
    cs = st.add_row().cells
    cell_text(cs[0],a); cell_text(cs[1],b); cell_text(cs[2],cval)

heading(doc, "How to use the real questionnaire")
for item in [
    "Copy Part A into Word, Google Forms, or Microsoft Forms and share it with genuine volunteer participants.",
    "Keep responses anonymous unless a participant explicitly agrees to provide contact details.",
    "Export genuine responses to Excel and create charts from the real data.",
    "Replace this synthetic appendix before reporting actual findings; retain it only as a chart-design example if needed.",
]:
    p = doc.add_paragraph(style='List Bullet'); p.paragraph_format.space_after = Pt(3); p.add_run(item)

doc.save(OUT)
print(OUT)
