from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = Path(__file__).resolve().parent / "VentureMind_Supervisory_Log_Sheets_5_to_7.docx"
PURPLE = "5B32D6"
NAVY = "111827"
LIGHT = "F5F3FF"


def set_cell(cell, fill=None):
    props = cell._tc.get_or_add_tcPr()
    if fill:
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), fill)
        props.append(shading)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def configure(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.12
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("VentureMind AI | CIS6035 Development Project | Supervisory Log Sheet")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(100, 116, 139)


def para(doc, text="", bold=False, size=10.5, color=None, align=None, before=0, after=5):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return p


def label(doc, text):
    return para(doc, text, bold=True, size=10.5, color=NAVY, before=7, after=3)


def metadata_table(doc, meeting, meeting_date):
    t = doc.add_table(rows=3, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    rows = [
        ("Student’s Name", "Vaishnavadevy Vasanthakumar", "University Number", "st20312021"),
        ("Date", meeting_date, "Meeting No.", meeting),
        ("Project Title", "VentureMind AI: AI-Powered Startup Lifecycle Management System", "Intake", "November 2025"),
    ]
    for r, values in enumerate(rows):
        c = t.rows[r].cells
        c[0].text = values[0]
        c[1].text = values[1]
        if r < 2:
            # add a nested two-column visual layout by using line breaks in the right cell
            c[1].text = values[1] + "\n\n" + values[2] + ": " + values[3]
        else:
            c[1].text = values[1] + "\n\n" + values[2] + ": " + values[3]
        set_cell(c[0], LIGHT)
        for cell in c:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9.5)
        c[0].paragraphs[0].runs[0].bold = True
    para(doc, "Supervisor’s Name: Mr. K. Sajuran     Supervisor Signature: ______________________________", size=9.5, after=1)
    para(doc, "Manager’s Name: T. Sutharsan        Manager’s Signature: ______________________________", size=9.5, after=5)


def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(item)
        r.font.size = Pt(10)


def sheet(doc, meeting, meeting_date, progression, discussion, action, final=False):
    para(doc, "ICBT", bold=True, size=18, color=PURPLE, align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
    para(doc, "International College of Business & Technology", bold=True, size=12, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER, after=6)
    para(doc, "ICBT Campus Project Log Sheet – Supervisory Sessions for CIS6035 Development Project", bold=True, size=13, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
    para(doc, "Note: Complete the supervisor and manager signature fields during the meeting. The suggested action list below must be confirmed or amended by the supervisor.", size=8.5, color="64748B", align=WD_ALIGN_PARAGRAPH.CENTER, after=7)
    metadata_table(doc, meeting, meeting_date)
    label(doc, "Work progression as to date (noted by student BEFORE mandatory supervisor meeting):")
    bullets(doc, progression)
    label(doc, "Items for Discussion (noted by student BEFORE mandatory supervisor meeting):")
    bullets(doc, discussion)
    label(doc, "Action List (to be attempted by student by the NEXT mandatory supervisory meeting – TO BE FILLED SUPERVISOR):")
    para(doc, "Suggested wording for supervisor confirmation:", size=9.5, bold=True, color=PURPLE, after=2)
    bullets(doc, action)
    if final:
        para(doc, "Final-session note: The student should attach final test evidence, the repository link/commit history, database schema evidence, signed declarations and any requested supervision records to the final submission pack.", size=9.5, color="7C2D12", before=4)
    para(doc, "Student signature: ______________________________     Date: __________________", size=9.5, before=8)


def main():
    doc = Document()
    configure(doc)
    sheet(doc, "05", "14.08.2026", [
        "Completed the initial integrated VentureMind AI prototype using React/Vite/TypeScript for the frontend and FastAPI, SQLAlchemy, Alembic and MySQL for the backend.",
        "Implemented role-based authentication and separate Founder, Administrator and Human Advisor workspaces.",
        "Developed the founder startup-profile workflow, explainable startup risk assessment, financial forecast, practical requirements checklist and Sri Lanka registration guidance.",
        "Implemented advisor directory, booking requests, availability, messaging, document-request flow and demonstration payment tracking.",
        "Implemented administrative areas for user management, advisor review, announcements, feedback, analytics and audit-oriented monitoring.",
        "Created early report-generation and launch/growth functions including social post text, poster design, hiring pack and first-month performance tracking."
    ], [
        "Review the user-interface design for a feasible, guided workflow rather than a collection of unrelated dashboard cards.",
        "Discuss effective backend integration, database migrations and reliable persistence of risk assessments, financial plans and bookings.",
        "Discuss Gemini API integration versus the local Ollama approach where an external API key is unavailable.",
        "Review chatbot behaviour, context-aware responses and safe structured fallback messages.",
        "Confirm the expected format and evidence for downloadable report creation, testing and final dissertation documentation."
    ], [
        "Refine the dashboard and workspace navigation so each founder sees a clear next step and role-appropriate actions.",
        "Complete end-to-end testing of backend endpoints and apply all outstanding Alembic migrations before the next review.",
        "Use Ollama local AI as the demonstrable default and document Gemini as an optional provider rather than relying on an unavailable API key.",
        "Resolve chatbot, report-generation and database persistence defects; capture screenshots and test evidence.",
        "Continue preparing the dissertation, including architecture diagrams, database design and a documented run guide."
    ])
    doc.add_page_break()
    sheet(doc, "06", "________________", [
        "Completed dashboard usability improvements, including clearer founder journey states, current-project details, next-action guidance and role-specific navigation.",
        "Integrated and tested explainable risk-analysis persistence, financial-planning workflow, registration-guide workflow and report-download components in the local environment.",
        "Configured local Ollama model support for context-aware AI guidance, while preserving a transparent deterministic fallback when the local service is unavailable.",
        "Extended advisor and administrator functionality: advisor profile/availability, booking-management views, notifications, user management, approvals, feedback and announcements.",
        "Implemented Launch & Growth Centre concepts: launch communications, editable poster/post content, hiring preparation, operations planning and first-month performance tracking.",
        "Prepared a full dissertation draft, diagrams, implementation explanations, testing matrix and selected interface evidence."
    ], [
        "Review whether the implemented scope satisfies the Development Project learning outcomes and identify the most important final refinements.",
        "Discuss evidence for testing: authentication, role-based access, saved profile, risk analysis, finance plan, booking, advisor messaging, reports and administration.",
        "Confirm how to present limitations honestly: live payment gateway, automatic social publishing, Google Places billing/API setup and optional external Gemini access.",
        "Review the dissertation structure, in-text citations, diagrams, screenshots, appendices and word count before final submission."
    ], [
        "Carry out final defect fixing and a full end-to-end demonstration with the backend, MySQL database and frontend running together.",
        "Prepare final screenshots, Swagger/API evidence, database schema evidence and test-result records for the dissertation appendix.",
        "Check that secrets are excluded from GitHub and that the repository README explains installation, migrations and local Ollama setup.",
        "Review the final report against the marking criteria, correct references/formatting and convert the final approved document to PDF.",
        "Prepare a short demonstration script covering Founder, Advisor and Administrator roles for the final supervisory session."
    ])
    doc.add_page_break()
    sheet(doc, "07", "________________", [
        "Completed the final integrated VentureMind AI prototype and checked the main Founder, Human Advisor and Administrator workflows in the local environment.",
        "Validated the core lifecycle path: registration/login, startup-profile completion, explainable risk analysis, financial plan, requirements, registration guidance, human support and launch/growth activities.",
        "Completed the final dissertation package containing objectives, literature review, methodology, requirements, architecture, database/UML diagrams, implementation, testing, evaluation, limitations and future work.",
        "Prepared the GitHub repository and local run instructions for the frontend, FastAPI backend, MySQL/Alembic migrations and optional Ollama local AI service.",
        "Recorded known limitations transparently: payment records are a demonstration workflow; social-media publishing requires future OAuth approval; Google Places and Gemini require separate API configuration; registration guidance does not replace official/legal processes."
    ], [
        "Obtain final supervisor feedback on the completed artefact, documentation quality and whether any high-priority corrections are required before submission.",
        "Confirm that the report follows the required format, word-count policy, Harvard referencing requirements and submission deadline.",
        "Confirm the final presentation/demo route and the evidence to submit with the source code and dissertation."
    ], [
        "Make only agreed final corrections to the artefact, report and appendices before submission.",
        "Export the approved dissertation to PDF, complete required declaration and signature pages, and submit the softcopy/hardcopy according to ICBT instructions.",
        "Retain a backup of the final PDF, source code, GitHub repository URL, database export and supervisory log sheets."
    ], final=True)
    doc.core_properties.title = "VentureMind AI Supervisory Log Sheets 5 to 7"
    doc.core_properties.author = "Vaishnavadevy Vasanthakumar"
    doc.save(OUT)
    print(f"Created {OUT}")


if __name__ == "__main__":
    main()
