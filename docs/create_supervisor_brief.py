from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUT = Path(__file__).with_name('VentureMind_AI_Supervisor_Brief.docx')

BLUE = '2E74B5'
DARK_BLUE = '1F4D78'
INK = '0B2545'
MUTED = '5B6573'
LIGHT = 'E8EEF5'
CALLOUT = 'F4F6F9'


def set_font(run, size=11, color='000000', bold=False, italic=False):
    run.font.name = 'Calibri'
    run._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
    run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in('w:tcMar')
    if tc_mar is None:
        tc_mar = OxmlElement('w:tcMar')
        tc_pr.append(tc_mar)
    for margin, value in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tc_mar.find(qn(f'w:{margin}'))
        if node is None:
            node = OxmlElement(f'w:{margin}')
            tc_mar.append(node)
        node.set(qn('w:w'), str(value))
        node.set(qn('w:type'), 'dxa')


def set_table_geometry(table, widths):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in('w:tblW')
    if tbl_w is None:
        tbl_w = OxmlElement('w:tblW')
        tbl_pr.append(tbl_w)
    tbl_w.set(qn('w:w'), str(sum(widths)))
    tbl_w.set(qn('w:type'), 'dxa')
    tbl_ind = OxmlElement('w:tblInd')
    tbl_ind.set(qn('w:w'), '120')
    tbl_ind.set(qn('w:type'), 'dxa')
    tbl_pr.append(tbl_ind)
    grid = table._tbl.tblGrid
    for grid_col, width in zip(grid.gridCol_lst, widths):
        grid_col.set(qn('w:w'), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in('w:tcW')
            if tc_w is None:
                tc_w = OxmlElement('w:tcW')
                tc_pr.append(tc_w)
            tc_w.set(qn('w:w'), str(width))
            tc_w.set(qn('w:type'), 'dxa')
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_para(p, before=0, after=6, line=1.25, align=None):
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    if align is not None:
        p.alignment = align


def add_text(doc, text, size=11, color='000000', bold=False, italic=False, before=0, after=6, line=1.25, align=None):
    p = doc.add_paragraph()
    set_para(p, before, after, line, align)
    set_font(p.add_run(text), size, color, bold, italic)
    return p


def heading(doc, text, level=1):
    size, color, before, after = {1: (16, BLUE, 18, 10), 2: (13, BLUE, 14, 7), 3: (12, DARK_BLUE, 10, 5)}[level]
    p = doc.add_paragraph()
    p.style = f'Heading {level}'
    set_para(p, before, after, 1.0)
    set_font(p.add_run(text), size, color, True)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    set_para(p, 0, 4, 1.25)
    set_font(p.add_run(text), 11)
    return p


def number(doc, text):
    p = doc.add_paragraph(style='List Number')
    set_para(p, 0, 4, 1.25)
    set_font(p.add_run(text), 11)
    return p


def add_callout(doc, label, body):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT)
    p = cell.paragraphs[0]
    set_para(p, 0, 0, 1.25)
    set_font(p.add_run(f'{label}: '), 10, DARK_BLUE, True)
    set_font(p.add_run(body), 10, INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    set_table_geometry(table, widths)
    for cell, value in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, LIGHT)
        p = cell.paragraphs[0]
        set_para(p, 0, 0, 1.1)
        set_font(p.add_run(value), 10, INK, True)
    for row in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            p = cell.paragraphs[0]
            set_para(p, 0, 0, 1.15)
            set_font(p.add_run(value), 9.5, '1F2937')
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    styles = doc.styles
    normal = styles['Normal']
    normal.font.name = 'Calibri'
    normal._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
    normal._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for level in [1, 2, 3]:
        style = styles[f'Heading {level}']
        style.font.name = 'Calibri'
        style._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
        style._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_para(header, 0, 0, 1.0)
    set_font(header.add_run('VENTUREMIND AI | SUPERVISOR BRIEF'), 8.5, MUTED, True)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para(footer, 0, 0, 1.0)
    set_font(footer.add_run('VentureMind AI - AI-Powered Startup Idea Evaluation and Business Planning Platform'), 8, MUTED)


def build():
    doc = Document()
    configure_document(doc)
    # Editorial cover
    add_text(doc, 'FINAL-YEAR SOFTWARE ENGINEERING PROJECT', 10, DARK_BLUE, True, after=16, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(doc, 'VentureMind AI', 30, INK, True, after=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(doc, 'Supervisor Brief, System Explanation, and Demonstration Workflow', 15, DARK_BLUE, False, after=30, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(doc, 'An explainable AI platform for startup idea evaluation and business planning', 11, MUTED, italic=True, after=60, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(doc, 'Prepared for project presentation', 11, INK, True, after=4, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(doc, 'July 2026', 10, MUTED, after=0, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()

    heading(doc, '1. Project Overview')
    add_text(doc, 'VentureMind AI is an AI-powered startup idea evaluation and business-planning platform. It helps entrepreneurs transform early ideas into clearer assumptions, evidence-focused questions, and practical next actions. The system does not promise that an idea will succeed. Instead, it supports better founder decisions through transparent scoring, structured learning resources, financial calculators, and an extensible full evaluation workflow.')
    add_callout(doc, 'Core principle', 'VentureMind is a decision-support platform. It makes assumptions visible and suggests what a founder should validate next.')
    heading(doc, 'What the platform currently offers', 2)
    for item in ['Startup Idea Generator for input-driven startup concepts.', 'Rapid Validation for a lightweight guest score, strength, and next step.', 'Ten browser-based founder calculators with displayed formulas.', 'Editable founder templates, checklists, stories, glossary, and book links.', 'Public pages for learning, tools, pricing, FAQ, contact, and project methodology.', 'Backend-ready architecture for authenticated projects, full evaluations, reports, chat, and administration.']:
        bullet(doc, item)

    heading(doc, '2. Main User Workflow')
    add_text(doc, 'The current public experience is designed for a guest founder who wants useful guidance without registration or a database connection.')
    for item in ['Visit VentureMind and choose Generate Ideas, Validate, Calculators, Templates, or Learning Resources.', 'Enter a startup idea or founder context.', 'Receive an immediate input-based result, concept, financial estimate, or structured template draft.', 'Use the output to plan customer interviews, experiments, pricing research, or a fuller project evaluation.', 'Move to the authenticated workspace when persistent projects, reports, and advanced AI analysis are enabled.']:
        number(doc, item)
    add_table(doc, ['Entry point', 'Input', 'Output'], [
        ('Startup Idea Generator', 'Industry, customer, skills, market, mode', 'Three startup hypotheses with revenue path, key risk, and MVP suggestion'),
        ('Rapid Validation', 'Idea description, optional industry and market', 'Indicative score, strength, and best next action'),
        ('Founder Calculators', 'Financial planning assumptions', 'Live financial estimate with visible formula'),
        ('Templates', 'Founder-provided fields', 'Copyable structured draft for notes or interviews'),
    ], [2200, 3300, 3860])

    heading(doc, '3. Rapid Validation Workflow')
    add_text(doc, 'Rapid Validation is the fastest guest feature. It is deliberately lightweight and clearly labelled as an indicative result rather than a market forecast.')
    for item in ['The user describes a startup idea and may add industry and market context.', 'The frontend checks explicit factors: description detail, industry, market, customer references, and problem or outcome references.', 'A deterministic score is calculated from those factors. No random score is generated.', 'The result shows one useful strength and one evidence-focused next step.']:
        number(doc, item)
    add_callout(doc, 'Important presentation point', 'Explain that the rapid score is an educational first-pass check. The full backend evaluation is the future workflow for detailed scoring, reports, and AI analysis.')

    heading(doc, '4. Founder Tool Suite')
    add_text(doc, 'The platform includes ten transparent calculators. All calculations run in the browser; visitor values are not stored by VentureMind in demo mode.')
    add_table(doc, ['Calculator group', 'Calculators', 'Purpose'], [
        ('Launch and cash', 'Startup Cost, Break-Even, Runway', 'Plan launch costs, sales threshold, and cash duration.'),
        ('Unit economics', 'ROI, CAC, LTV', 'Measure spend efficiency and customer value assumptions.'),
        ('Market and funding', 'Market Size, Funding, Equity Dilution, Startup Valuation', 'Estimate TAM/SAM/SOM, raise target, ownership, and valuation scenarios.'),
    ], [1800, 3600, 3960])
    heading(doc, 'Example formula', 3)
    add_text(doc, 'Break-even units = Fixed monthly costs / (Average sale price - Variable cost per sale)', 11, DARK_BLUE, True, before=0, after=8)
    add_text(doc, 'Every calculator displays its formula on-screen so the user can understand the result and change the assumptions.')

    heading(doc, '5. Learning and Founder Resources')
    add_table(doc, ['Resource', 'Purpose'], [
        ('Founder Templates', 'Customer Persona, Validation Interview, Competitor Analysis, Problem Hypothesis, Experiment Plan, and One-Minute Pitch.'),
        ('Founder Checklist', 'Interactive first-pass checklist for customer, problem, alternative, assumptions, and experiments.'),
        ('Startup Stories', 'Source-linked success and failure examples, with documented facts separated from the learning takeaway.'),
        ('Startup Books', 'Links to legitimate author, publisher, or official book pages, including The Lean Startup, The Mom Test, Venture Deals, and INSPIRED.'),
        ('Explainable AI', 'A plain-language explanation of visible factors, deterministic foundations, and human responsibility.'),
    ], [2200, 7160])

    heading(doc, '6. Full-System Architecture')
    add_text(doc, 'The application follows a modular full-stack architecture. The public demo features can run without a backend, while the production architecture supports persistent user projects and advanced evaluation.')
    add_table(doc, ['Layer', 'Technology', 'Responsibility'], [
        ('Frontend', 'React, Vite, TypeScript, Tailwind CSS, React Router', 'Responsive pages, forms, calculators, templates, dashboards, and user interactions.'),
        ('API', 'FastAPI, Pydantic, JWT', 'REST APIs, validation, security boundaries, and consistent responses.'),
        ('Business services', 'Python service layer', 'Project logic, deterministic evaluation, report generation, chat context, and admin services.'),
        ('Data', 'SQLAlchemy, Alembic, MySQL', 'Normalized persistence for users, projects, ideas, evaluations, reports, feedback, and chat history.'),
        ('AI integration', 'Gemini-compatible LLM interface', 'Business analysis and context-aware assistant capabilities when configured.'),
    ], [1500, 3300, 4560])

    heading(doc, '7. Complete Evaluation Workflow')
    add_text(doc, 'When the backend and MySQL are enabled, the full evaluation workflow is:')
    for item in ['Authenticated user creates a project and submits structured startup information.', 'FastAPI validates the request and stores project and idea data.', 'The NLP layer extracts industry, customer, problem, solution, business model, and other structured information.', 'The evaluation engine calculates documented weighted scores.', 'The explainability layer identifies positive factors, negative factors, missing assumptions, and improvement suggestions.', 'The LLM layer generates business analysis, recommendations, roadmap, and contextual assistant responses.', 'The system stores the evaluation, displays results, and can generate a downloadable PDF report.']:
        number(doc, item)

    heading(doc, '8. Explainable AI Design')
    add_text(doc, 'Explainability is a central project requirement. A score without a reason is not useful for a founder. Each full evaluation dimension is designed to show the score, the evidence affecting it, the missing evidence, risk factors, and a practical improvement action.')
    add_table(doc, ['Example evaluation dimension', 'What the user sees'], [
        ('Market Opportunity', 'Score, identified target customer, evidence gaps, and suggested market interviews.'),
        ('Business Model', 'Revenue model clarity, pricing assumptions, gaps, and a validation recommendation.'),
        ('Technical Feasibility', 'Technical scope assumptions, dependencies, risks, and an MVP recommendation.'),
        ('Investment Readiness', 'Readiness factors, missing evidence, and practical next actions before fundraising.'),
    ], [2800, 6560])
    add_callout(doc, 'Supervisor explanation', 'The platform uses deterministic evaluation factors for core scoring and uses generative AI for narrative recommendations. This separates measurable logic from generated language.')

    heading(doc, '9. Data Model and Roles')
    add_table(doc, ['Role', 'Access'], [
        ('Guest', 'Public pages, idea generation, rapid validation, calculators, templates, checklists, stories, and books.'),
        ('User', 'Projects, stored startup ideas, full evaluations, reports, chat history, and notifications.'),
        ('Admin', 'User management, analytics, AI usage review, and feedback administration.'),
    ], [1800, 7560])
    add_text(doc, 'The normalized database design includes Users, Security Tokens, Projects, Startup Ideas, Evaluations, Evaluation Scores, Reports, Feedback, Notifications, Chat Conversations, and Chat Messages.')

    heading(doc, '10. Demo Mode and Production Mode')
    add_table(doc, ['Demo mode now', 'Production mode later'], [
        ('Guest-first frontend tools run without login, MySQL, or backend connection.', 'MySQL, authentication, persistent projects, email delivery, full AI analysis, reports, and admin data are enabled.'),
        ('Rapid Validation and calculators are intentionally local and transparent.', 'Full evaluation uses backend services, stored data, explainability output, and LLM-generated analysis.'),
    ], [4680, 4680])
    add_text(doc, 'This distinction is important in the presentation: demo mode makes the system easy to demonstrate locally, while the backend architecture shows how the final deployed platform will scale.')

    heading(doc, '11. Suggested Live Demonstration Script')
    for item in ['Open the landing page and explain that VentureMind supports founders before they invest major time or money.', 'Use Rapid Validation with a sample startup idea. Explain that the result is deterministic, fast, and indicative.', 'Use the Idea Generator with an industry, customer type, skills, and market. Show the generated hypothesis and MVP suggestion.', 'Open one or two calculators, change values, and point out the displayed formula.', 'Open a template and show how a founder can create a customer interview or problem hypothesis draft.', 'Show Startup Stories and Books to demonstrate the learning ecosystem.', 'Explain that a logged-in user will later move into the full backend evaluation and PDF-report workflow.']:
        number(doc, item)

    heading(doc, '12. Likely Supervisor Questions')
    add_table(doc, ['Question', 'Recommended answer'], [
        ('Can the platform predict startup success?', 'No. It is a decision-support tool. It helps founders identify assumptions, evidence gaps, and next actions.'),
        ('Why is the rapid score not random?', 'It uses explicit input factors such as idea detail, industry, market, customer references, and problem references.'),
        ('What makes the AI explainable?', 'The full design exposes positive factors, negative factors, missing assumptions, suggestions, and score reasoning.'),
        ('Why do you have guest mode?', 'It reduces friction and makes the platform demonstrable without database setup. Persistent projects remain part of production mode.'),
        ('What is left before production deployment?', 'Configure MySQL and environment variables, enable authentication and email delivery, connect the LLM, complete tests, and deploy the API and frontend.'),
    ], [2600, 6760])

    heading(doc, '13. Key Files to Show in VS Code')
    for item in ['frontend/src/routes/AppRouter.tsx - all public routes and workspace routes.', 'frontend/src/pages/ValidateLandingPage.tsx - rapid guest validation.', 'frontend/src/pages/CalculatorsPage.tsx - transparent financial formulas.', 'frontend/src/pages/TemplatesPage.tsx - editable founder templates.', 'frontend/src/pages/StartupStoriesPage.tsx and StartupBooksPage.tsx - curated learning resources.', 'backend/app - FastAPI application, data models, APIs, evaluation services, reports, chat, and admin modules.']:
        bullet(doc, item)

    heading(doc, '14. Presentation Closing')
    add_text(doc, 'VentureMind AI combines practical founder tools with a scalable backend architecture for explainable startup evaluation. Its central value is not a promise of success; it is helping founders ask better questions, test important assumptions earlier, and make clearer next decisions.')
    doc.core_properties.title = 'VentureMind AI Supervisor Brief'
    doc.core_properties.subject = 'Project explanation, workflows, and presentation notes'
    doc.core_properties.author = 'VentureMind AI Project Team'
    doc.save(OUT)


if __name__ == '__main__':
    build()
    print(OUT)
